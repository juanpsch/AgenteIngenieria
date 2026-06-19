"""Lectura de documentos — híbrido texto + visión (T0.8).

Dado un archivo, devuelve una representación legible para los agentes:
- `contenido`: texto/markdown extraído (lo que el LLM lee).
- `imagenes`: data-URLs base64 PNG de páginas/planos sin texto (para visión).

Router por extensión. Cada lector está aislado en try/except: si uno falla, el
documento se marca `legible=False` con `motivo`, **sin** tumbar el flujo.

Stack: PyMuPDF (PDF texto+render), pdfplumber (tablas), openpyxl (xlsx),
python-docx (docx), ezdxf (DXF) + ezdxf.addons.odafc (DWG, requiere ODA Converter).
"""

from __future__ import annotations

import base64
import os
from pathlib import Path
from typing import Any

# Límites para no inflar el payload del LLM
MAX_TEXT_CHARS = 20_000
MAX_PAGES_RENDER = 5
RENDER_DPI = 120


def _result(filename: str, ext: str, **kw: Any) -> dict[str, Any]:
    base = {
        "filename": filename,
        "ext": ext,
        "legible": False,
        "contenido": "",
        "imagenes": [],
        "motivo": "",
    }
    base.update(kw)
    if base["contenido"] and len(base["contenido"]) > MAX_TEXT_CHARS:
        base["contenido"] = base["contenido"][:MAX_TEXT_CHARS] + "\n…[truncado]"
    return base


def _png_data_url(png_bytes: bytes) -> str:
    return "data:image/png;base64," + base64.b64encode(png_bytes).decode("ascii")


def render_pdf_images(path: str, max_pages: int = 4, dpi: int = 110) -> list[str]:
    """Renderiza las primeras páginas de un PDF a data-URLs PNG.

    Para CAPTURA de templates: a diferencia de read_document (que solo renderiza
    páginas sin texto), acá rendereamos siempre, así el extractor 've' el cajetín,
    el índice de hojas y los diagramas aunque el PDF tenga texto.
    """
    p = Path(path)
    if p.suffix.lower().lstrip(".") != "pdf" or not p.exists():
        return []
    try:
        import fitz

        out: list[str] = []
        with fitz.open(str(p)) as doc:
            for pno in range(min(doc.page_count, max_pages)):
                pix = doc[pno].get_pixmap(dpi=dpi)
                out.append(_png_data_url(pix.tobytes("png")))
        return out
    except Exception:
        return []


def contar_paginas(path: str) -> int:
    """Nº de páginas de un PDF (1 para imágenes / no-PDF). 0 si no se puede abrir."""
    p = Path(path)
    if not p.exists():
        return 0
    if p.suffix.lower().lstrip(".") != "pdf":
        return 1
    try:
        import fitz

        with fitz.open(str(p)) as doc:
            return int(doc.page_count)
    except Exception:
        return 0


def buscar_texto(path: str, q: str) -> list[dict[str, Any]]:
    """Busca `q` en TODO el PDF (capa de texto, PyMuPDF). Devuelve [{pagina(1-idx), rects:[{x,y,w,h}]}]
    con bboxes relativos (0–1) para resaltar. [] si no es PDF, no hay texto, o no se encuentra."""
    q = (q or "").strip()
    p = Path(path)
    if not q or p.suffix.lower().lstrip(".") != "pdf" or not p.exists():
        return []
    try:
        import fitz

        out: list[dict[str, Any]] = []
        with fitz.open(str(p)) as doc:
            for i, page in enumerate(doc):
                w = float(page.rect.width) or 1.0
                h = float(page.rect.height) or 1.0
                rects = page.search_for(q) or []
                if rects:
                    out.append({"pagina": i + 1, "rects": [
                        {"x": round(r.x0 / w, 4), "y": round(r.y0 / h, 4),
                         "w": round((r.x1 - r.x0) / w, 4), "h": round((r.y1 - r.y0) / h, 4)} for r in rects]})
        return out
    except Exception:
        return []


def _pagina_vacia(page) -> bool:
    """Página sin contenido: sin texto, sin imágenes y sin vectores. Barato (corta apenas encuentra algo)."""
    try:
        if page.get_text("text").strip():
            return False
        if page.get_images():
            return False
        if page.get_drawings():
            return False
        return True
    except Exception:
        return False  # ante la duda, NO la consideramos vacía


def paginas_utiles(path: str) -> int:
    """Nº de páginas hasta la ÚLTIMA con contenido (recorta las hojas EN BLANCO del final) para que la
    previsualización no muestre páginas vacías de más. Escanea desde el final hacia atrás (barato:
    solo toca las hojas finales). Cae a `contar_paginas` si no es PDF o todo está vacío."""
    p = Path(path)
    if p.suffix.lower().lstrip(".") != "pdf" or not p.exists():
        return contar_paginas(path)
    try:
        import fitz

        with fitz.open(str(p)) as doc:
            for i in range(doc.page_count - 1, -1, -1):
                if not _pagina_vacia(doc[i]):
                    return i + 1
            return int(doc.page_count)  # todas vacías -> no recortar
    except Exception:
        return contar_paginas(path)


def render_pdf_page(path: str, page: int = 1, dpi: int = 110) -> str | None:
    """Renderiza UNA página puntual (1-indexada) de un PDF a data-URL PNG. None si no existe."""
    p = Path(path)
    if p.suffix.lower().lstrip(".") != "pdf" or not p.exists():
        return None
    try:
        import fitz

        with fitz.open(str(p)) as doc:
            idx = max(0, page - 1)
            if idx >= doc.page_count:
                return None
            return _png_data_url(doc[idx].get_pixmap(dpi=dpi).tobytes("png"))
    except Exception:
        return None


def read_document(path: str) -> dict[str, Any]:
    """Lee un archivo y devuelve {filename, ext, legible, contenido, imagenes, motivo}."""
    p = Path(path)
    ext = p.suffix.lower().lstrip(".")
    filename = p.name

    if not p.exists():
        return _result(filename, ext, motivo="archivo no encontrado")

    try:
        if ext == "pdf":
            return _read_pdf(p, filename, ext)
        if ext == "xlsx":
            return _read_xlsx(p, filename, ext)
        if ext == "docx":
            return _read_docx(p, filename, ext)
        if ext == "dxf":
            return _read_dxf(p, filename, ext)
        if ext == "dwg":
            return _read_dwg(p, filename, ext)
        if ext in ("png", "jpg", "jpeg", "webp", "bmp", "tif", "tiff"):
            return _read_image(p, filename, ext)
        if ext in ("txt", "md", "csv"):
            return _result(filename, ext, legible=True,
                           contenido=p.read_text(encoding="utf-8", errors="replace"))
    except Exception as exc:  # un lector que falla no rompe el flujo
        return _result(filename, ext, motivo=f"error leyendo {ext}: {exc}")

    return _result(filename, ext, motivo=f"formato no soportado: .{ext}")


# --- Lectores por formato ----------------------------------------------------

def _read_pdf(p: Path, filename: str, ext: str) -> dict[str, Any]:
    import fitz  # PyMuPDF

    textos: list[str] = []
    imagenes: list[str] = []
    with fitz.open(str(p)) as doc:
        for page in doc:
            t = page.get_text("text")
            if t and t.strip():
                textos.append(t)
            elif len(imagenes) < MAX_PAGES_RENDER:
                pix = page.get_pixmap(dpi=RENDER_DPI)
                imagenes.append(_png_data_url(pix.tobytes("png")))

    # Tablas (best-effort, no crítico)
    tablas_md = _pdf_tables(p)

    contenido = "\n\n".join(textos)
    if tablas_md:
        contenido += "\n\n## Tablas detectadas\n" + tablas_md
    if not contenido.strip() and imagenes:
        # PDF sin capa de texto (escaneo): OCR de las páginas (si hay) para captura/cobertura/extracción.
        try:
            from tools import ocr

            if ocr.disponible():
                ocrtxt = "\n\n".join(filter(None, (ocr.ocr_texto(im) for im in imagenes)))
                if ocrtxt.strip():
                    contenido = ocrtxt
        except Exception:
            pass
    if not contenido.strip() and imagenes:
        contenido = "[PDF sin texto extraíble — se adjuntan páginas como imagen para visión]"

    return _result(filename, ext, legible=True, contenido=contenido, imagenes=imagenes)


def _pdf_tables(p: Path) -> str:
    try:
        import pdfplumber

        out: list[str] = []
        with pdfplumber.open(str(p)) as pdf:
            for i, page in enumerate(pdf.pages[:MAX_PAGES_RENDER]):
                for table in page.extract_tables() or []:
                    rows = ["| " + " | ".join((c or "") for c in row) + " |" for row in table]
                    if rows:
                        out.append(f"(página {i + 1})\n" + "\n".join(rows))
        return "\n\n".join(out)
    except Exception:
        return ""


def _read_xlsx(p: Path, filename: str, ext: str) -> dict[str, Any]:
    import openpyxl

    wb = openpyxl.load_workbook(str(p), read_only=True, data_only=True)
    partes: list[str] = []
    for ws in wb.worksheets:
        partes.append(f"### Hoja: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            celdas = [str(c) if c is not None else "" for c in row]
            if any(celdas):
                partes.append("| " + " | ".join(celdas) + " |")
    wb.close()
    return _result(filename, ext, legible=True, contenido="\n".join(partes))


def _read_docx(p: Path, filename: str, ext: str) -> dict[str, Any]:
    import docx

    d = docx.Document(str(p))
    partes = [par.text for par in d.paragraphs if par.text.strip()]
    for table in d.tables:
        for row in table.rows:
            partes.append("| " + " | ".join(cell.text for cell in row.cells) + " |")
    return _result(filename, ext, legible=True, contenido="\n".join(partes))


def _dxf_to_text(doc: Any) -> str:
    partes: list[str] = []
    try:
        partes.append("Capas: " + ", ".join(layer.dxf.name for layer in doc.layers))
    except Exception:
        pass
    msp = doc.modelspace()
    textos: list[str] = []
    for e in msp:
        try:
            if e.dxftype() == "TEXT":
                textos.append(e.dxf.text)
            elif e.dxftype() == "MTEXT":
                textos.append(e.text)
        except Exception:
            continue
    if textos:
        partes.append("Textos del dibujo:\n" + "\n".join(textos))
    return "\n".join(partes) if partes else "[DXF sin textos legibles]"


def _read_dxf(p: Path, filename: str, ext: str) -> dict[str, Any]:
    import ezdxf

    doc = ezdxf.readfile(str(p))
    return _result(filename, ext, legible=True, contenido=_dxf_to_text(doc))


def _read_dwg(p: Path, filename: str, ext: str) -> dict[str, Any]:
    # DWG es propietario: se convierte a DXF con ODA File Converter vía ezdxf.
    try:
        from ezdxf.addons import odafc

        oda_path = os.getenv("ODA_CONVERTER_PATH")
        if oda_path:
            odafc.win_exec_path = oda_path
        doc = odafc.readfile(str(p))
        return _result(filename, ext, legible=True, contenido=_dxf_to_text(doc))
    except Exception as exc:
        return _result(
            filename, ext, legible=False,
            motivo=(
                "DWG no legible: requiere ODA File Converter instalado "
                f"(ODA_CONVERTER_PATH). Detalle: {exc}"
            ),
        )


def _read_image(p: Path, filename: str, ext: str) -> dict[str, Any]:
    raw = p.read_bytes()
    mime = "image/jpeg" if ext in ("jpg", "jpeg") else f"image/{ext}"
    data_url = f"data:{mime};base64," + base64.b64encode(raw).decode("ascii")
    return _result(
        filename, ext, legible=True,
        contenido="[imagen — se adjunta para análisis por visión]",
        imagenes=[data_url],
    )
